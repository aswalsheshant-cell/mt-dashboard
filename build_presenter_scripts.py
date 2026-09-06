"""Build MT_July26_Presenter_Scripts.docx — Slide # | Title | Script table."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPTS = [
    {
        "slide": 1,
        "title": "Cover — MT Performance & Market Share Review",
        "script": (
            "Welcome to the Modern Trade Performance and Market Share Review for July 2026, "
            "presented in the context of Q1 FY27 performance. This deck covers three channels: "
            "MT Zones, eB2B — our Nykaa/FSN business — and SIS, Shop-in-Shop. "
            "One important governance note before we begin: June offtake source data is absent. "
            "This is a Tier 3 pipeline block under our three-tier pre-flight protocol — meaning "
            "we cannot compute a complete Q1 total. All numbers you see are exact from verified "
            "source files; no estimates or fabricated figures have been used."
        ),
    },
    {
        "slide": 2,
        "title": "Q1 FY27 Offtake Snapshot · Apr – Jun '26",
        "script": (
            "Starting with Q1 FY27 context as requested. April delivered ₹33.60 Crore NSV on "
            "19.6 lakh units — a strong opening to the financial year. May grew 13% month-on-month "
            "to ₹38.11 Crore on 22 lakh units, our highest month in the period covered. "
            "June is a Tier 3 block: the source file is absent from our data folders, so we cannot "
            "report that month and Q1 is partial, covering only April and May — ₹71.71 Crore combined. "
            "July at ₹33.96 Crore opens Q2. "
            "The key insight to carry is the ASP premiumisation story: average selling price moved "
            "from ₹171.7 in April to ₹179.2 in July — a 4.4% appreciation in four months — "
            "while realisation improved from 41.6% to 42.2%. The portfolio is moving up in value."
        ),
    },
    {
        "slide": 3,
        "title": "July '26 · MT Channel Overview",
        "script": (
            "Now to July's headline numbers. MT primary at ₹47.02 Crore; MT offtake at ₹33.96 Crore. "
            "Conversion is 72.2% — that is 13.5 percentage points below our West and South-1 benchmark "
            "of 85.7%. The gap between primary and offtake is ₹13.06 Crore. "
            "Not all of that gap is recoverable — some is structural timing or floor stock. "
            "The actionable portion: ₹6.22 Crore sits above our ₹0.25 Crore per-zone floor. "
            "That is the number the field team should be chasing this month. "
            "Separately, eB2B through Nykaa: ₹2.20 Crore primary, ₹2.07 Crore offtake, "
            "93.9% flow — the best-converting channel we have. "
            "SIS is very small and net of MRN returns. "
            "Neither eB2B nor SIS is included in zone figures — they are separate channels per our governance."
        ),
    },
    {
        "slide": 4,
        "title": "Zone Conversion Analysis · July '26",
        "script": (
            "Breaking the conversion picture down by zone. "
            "South-1 leads at 86.3% and West follows at 85.2% — both at or above the 85.7% benchmark. "
            "These are our model geographies; the practices driving their conversion should be documented "
            "and replicated. "
            "Central is next at 80.9% — reasonable but with room. "
            "South-2 at 72.4% is borderline. "
            "North at 61.3% has the largest absolute gap at ₹4.41 Crore — this is a volume issue. "
            "East is the most critical at 49.9%: less than half the primary stock has converted to "
            "shelf offtake. With ₹3.56 Crore gap, East needs an immediate field investigation. "
            "In absolute terms, North and East account for ₹7.97 Crore of the ₹13.06 Crore total gap — "
            "61% of the problem is in two zones."
        ),
    },
    {
        "slide": 5,
        "title": "Account Gap Analysis · July '26",
        "script": (
            "Zooming into accounts. Reliance is the critical gap account: ₹7.61 Crore gap "
            "at 51.5% conversion. That is 58% of the entire MT gap in one account. "
            "Moving Reliance to the 75% level would recover approximately ₹3.7 Crore — "
            "that alone changes the conversion picture materially. "
            "DMart at ₹4.29 Crore gap is likely a stocking lag — they tend to build inventory "
            "ahead of billing cycles, so this may self-correct. Monitor closely. "
            "Metro at ₹1.36 Crore is manageable. "
            "On the positive side: Apollo is our gold standard account at 99.7% conversion — "
            "₹7.20 Crore primary, ₹7.18 Crore offtake. Whatever Apollo's account manager is doing, "
            "capture it. "
            "Lulu is a special case: zero primary billing but ₹1.70 Crore offtake. "
            "This is a consignment or billing lag pattern — stores are drawing from existing stock "
            "without fresh billing. Not necessarily a problem but needs monitoring to ensure "
            "replenishment happens before stockouts. "
            "The recoverable gap above floor is ₹6.22 Crore — this is the addressable number "
            "for the next 30 days."
        ),
    },
    {
        "slide": 6,
        "title": "Market Share · Nielsen RMS",
        "script": (
            "Market share from Nielsen RMS. Two categories to focus on. "
            "In Face Wash within MT, Mamaearth is ranked fourth at 10.5%. "
            "Himalaya leads at 22.6%, Garnier at 14.2%, Pond's at 13.8%. "
            "We are competitive but there is a 12-point gap to the category leader. "
            "The sharper story is Shampoo: Mamaearth at 3.7% versus Dove at 16.6% and H&S at 13.0%. "
            "That is a 4-times gap to the category leader — shampoo is our biggest MT market share opportunity. "
            "One governance note: the July Nielsen cut is pending. A Tier 3 badge is applied — "
            "these numbers represent the latest available cut, which is June or earlier. "
            "When the July data comes in, we should re-run this slide."
        ),
    },
    {
        "slide": 7,
        "title": "Brand Primary vs Offtake · July '26",
        "script": (
            "Brand-level P vs O for July. Mamaearth is our volume engine: "
            "₹33.38 Crore primary, ₹24.49 Crore offtake — 73.4% conversion, ₹8.89 Crore gap. "
            "The Derma Co. at ₹15.19 Crore primary, ₹11.03 Crore offtake — 72.6% conversion, "
            "₹4.16 Crore gap. Both brands track close together in conversion rate. "
            "Aqualogica shows offtake slightly above primary at ₹0.48 versus ₹0.41 Crore — "
            "a stock draw pattern, likely from prior month inventory. "
            "BBlunt and Dr. Sheth's are very small volumes — keep monitoring but not a focus area today. "
            "Combined, the portfolio has meaningful conversion headroom."
        ),
    },
    {
        "slide": 8,
        "title": "Account Primary vs Offtake · July '26",
        "script": (
            "Account P vs O gives us the clearest view of where conversion is happening and where it is not. "
            "DMart is our largest primary account at ₹18.25 Crore but offtake at ₹13.97 Crore, "
            "76.5% — the stocking lag pattern. Reliance at ₹15.66 Crore primary, "
            "₹8.06 Crore offtake, 51.5% — this is where the most urgent action is needed. "
            "Apollo at ₹7.20 Crore primary, ₹7.18 Crore offtake, 99.7% — the gold standard. "
            "FSN/Nykaa at 99.5% reflects the eB2B channel's strong flow — this is included here "
            "for brand-level visibility but is a separate channel in the zone figures. "
            "Lulu: zero primary, ₹1.70 Crore offtake — consignment draw. "
            "Wellness Forward and H&G both show offtake exceeding primary — similar consignment patterns. "
            "Metro at 26.6% conversion is the weakest account conversion — only ₹0.49 Crore offtake "
            "against ₹1.84 Crore primary; this needs investigation."
        ),
    },
    {
        "slide": 9,
        "title": "Category Offtake Trends · Feb – Jul '26",
        "script": (
            "Category trend analysis over the Feb to July window. Starting with Mamaearth. "
            "Face Cleanser peaked in May and June at ₹9.65 Crore and pulled back to ₹8.53 Crore "
            "in July — likely a seasonal normalisation after summer peak demand. "
            "Shampoo is the standout story: from ₹4.81 Crore in February to ₹6.95 Crore in July, "
            "up 44.5% in six months, showing sustained momentum without a seasonal spike. "
            "Sun Care is retreating as expected post-summer — from ₹3.10 Crore in April "
            "to ₹1.30 Crore in July. "
            "Now The Derma Co.: Face Cleanser jumped from ₹4.83 Crore in June to ₹7.13 Crore in July — "
            "a 47.5% month-on-month spike. This warrants immediate investigation: "
            "is this a large account listing, a promotional fill, a channel event? "
            "Understanding the driver will tell us if this is repeatable or one-off. "
            "Sun Care and Face Serum for Derma Co. are relatively stable. "
            "On eB2B through Nykaa: the April spike to ₹2.29 Crore was notable; "
            "July settled at ₹2.07 Crore. Active EANs have declined from 222 in January to 198 in July — "
            "this suggests deliberate portfolio pruning, which should improve velocity on the "
            "remaining SKUs. Watch if this leads to better per-EAN offtake in the coming months."
        ),
    },
]

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def main():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # Title
    title = doc.add_heading('MT Performance Review — July 2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x0D, 0x1E, 0x35)
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph('Presenter Scripts by Slide  ·  Honasa Consumer Ltd. · MT Analytics')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.runs[0]
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x8A, 0x9B, 0xB2)
    sr.font.italic = True

    doc.add_paragraph('')

    # Note box
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        'Usage: These scripts are intended for the presenter only and do NOT appear in the deck. '
        'Read naturally — adjust to your audience. Data governance notes within scripts are '
        'for presenter awareness; share selectively.'
    )
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(0x6B, 0x5B, 0x00)
    note_para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph('')

    # Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    widths = [Cm(1.8), Cm(5.0), Cm(11.2)]
    hdr_texts = ['Slide #', 'Slide Title', 'Presenter Script']
    hdr_fills = ['0D1E35', '0D1E35', '0D1E35']
    for i, cell in enumerate(hdr.cells):
        cell.width = widths[i]
        set_cell_bg(cell, hdr_fills[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr_texts[i])
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Script rows
    row_fills = ['FFFFFF', 'F4F6FB']
    for idx, item in enumerate(SCRIPTS):
        row = table.add_row()
        row_bg = row_fills[idx % 2]

        # Slide number
        c0 = row.cells[0]
        c0.width = widths[0]
        set_cell_bg(c0, row_bg)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(item['slide']))
        r0.bold = True
        r0.font.size = Pt(12)
        r0.font.name = 'Cambria'
        r0.font.color.rgb = RGBColor(0x0D, 0x1E, 0x35)

        # Title
        c1 = row.cells[1]
        c1.width = widths[1]
        set_cell_bg(c1, row_bg)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(item['title'])
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = 'Cambria'
        r1.font.color.rgb = RGBColor(0x0D, 0x1E, 0x35)

        # Script
        c2 = row.cells[2]
        c2.width = widths[2]
        set_cell_bg(c2, row_bg)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(item['script'])
        r2.font.size = Pt(9.5)
        r2.font.name = 'Calibri'
        r2.font.color.rgb = RGBColor(0x1A, 0x28, 0x40)
        p2.paragraph_format.space_after = Pt(0)

    # Footer
    doc.add_paragraph('')
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer_para.add_run(
        'Honasa Consumer Ltd. · MT Analytics · August 2026  |  '
        'Tier 3 note: Jun offtake source absent; Q1 partial (Apr+May). '
        'Nielsen Jul MS pending.'
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x8A, 0x9B, 0xB2)
    fr.font.italic = True

    out = '/home/user/mt-dashboard/MT_July26_Presenter_Scripts.docx'
    doc.save(out)
    print(f'DONE: {out}')

if __name__ == '__main__':
    main()
